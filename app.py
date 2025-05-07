import streamlit as st
from docx import Document
from PyPDF2 import PdfReader
import os
import io

def run_length_encode_custom(text, marker="#"):
    if not text:
        return ""
    encoded = ""
    prev_char = text[0]
    count = 1

    for char in text[1:]:
        if char == prev_char:
            count += 1
        else:
            if count >= 3:
                encoded += f"{marker}{count}{prev_char}"
            else:
                encoded += prev_char * count
            prev_char = char
            count = 1

    if count >= 3:
        encoded += f"{marker}{count}{prev_char}"
    else:
        encoded += prev_char * count

    return encoded

def run_length_decode_custom(text, marker="#"):
    decoded = ""
    i = 0
    while i < len(text):
        if text[i] == marker:
            i += 1
            num = ""
            while i < len(text) and text[i].isdigit():
                num += text[i]
                i += 1
            if i < len(text):
                decoded += text[i] * int(num)
                i += 1
        else:
            decoded += text[i]
            i += 1
    return decoded

# -------------------- File Processing --------------------

def extract_text_from_docx(file) -> str:
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def create_docx_from_text(text: str) -> io.BytesIO:
    doc = Document()
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_text_from_pdf(file) -> str:
    reader = PdfReader(file)
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def create_txt_buffer(content: str) -> io.BytesIO:
    buffer = io.BytesIO()
    buffer.write(content.encode('utf-8'))
    buffer.seek(0)
    return buffer

def get_size_in_kb(size_bytes):
    return round(size_bytes / 1024, 2)

# -------------------- UI --------------------

st.set_page_config(page_title="SiKompres", page_icon="📄", layout="centered")

st.markdown("""
    <h1 style='text-align: center; color: #4CAF50;'>📄 My Kompres</h1>
    <p style='text-align: center;'>🔧 Kompresi & Dekompresi File menggunakan <b>Run-Length Encoding</b> (RLE)</p>
    """, unsafe_allow_html=True)

with st.container():
    st.markdown("### 🔽 Pilih Mode Operasi")
    mode = st.radio("Apa yang ingin Anda lakukan?", ["🗜️ Kompresi", "📂 Dekompresi"], horizontal=True, label_visibility="collapsed")
    mode = "Kompresi" if "Kompresi" in mode else "Dekompresi"

    st.markdown("### 📁 Unggah File Anda")
    uploaded_file = st.file_uploader(
        "Pilih file berformat `.docx`, `.txt`, atau `.pdf`", 
        type=["docx", "txt", "pdf"],
        help="File PDF hanya bisa dikompresi, bukan didekompresi."
    )

# -------------------- File Handling --------------------

if uploaded_file:
    file_bytes = uploaded_file.read()
    filename, ext = os.path.splitext(uploaded_file.name)
    result_io = io.BytesIO()
    result_ext = ext

    with st.spinner("⏳ Memproses file..."):
        progress = st.progress(10)
        original_size = get_size_in_kb(len(file_bytes))
        progress.progress(30)

        if mode == "Kompresi":
            if ext == ".docx":
                text = extract_text_from_docx(io.BytesIO(file_bytes))
                encoded = run_length_encode_custom(text)
                result_io = create_docx_from_text(encoded)

            elif ext == ".txt":
                text = file_bytes.decode("utf-8", errors="ignore")
                encoded = run_length_encode_custom(text)
                result_io = create_txt_buffer(encoded)

            elif ext == ".pdf":
                text = extract_text_from_pdf(io.BytesIO(file_bytes))
                encoded = run_length_encode_custom(text)
                result_io = create_txt_buffer(encoded)
                result_ext = ".txt"

        else:  # Dekompresi
            if ext == ".docx":
                text = extract_text_from_docx(io.BytesIO(file_bytes))
                decoded = run_length_decode_custom(text)
                result_io = create_docx_from_text(decoded)

            elif ext == ".txt":
                text = file_bytes.decode("utf-8", errors="ignore")
                decoded = run_length_decode_custom(text)
                result_io = create_txt_buffer(decoded)

            elif ext == ".pdf":
                st.warning("⚠️ File PDF tidak bisa didekompresi secara langsung. Gunakan file hasil kompresi (.txt/.docx).")
                result_io = None

        progress.progress(80)

        if result_io:
            result_size = get_size_in_kb(len(result_io.getvalue()))
            download_name = filename + ("_compressed" if mode == "Kompresi" else "_decompressed") + result_ext

            with st.expander("📊 Lihat Ringkasan Ukuran File"):
                col1, col2 = st.columns(2)
                col1.metric("📄 Ukuran Asli", f"{original_size} KB")
                col2.metric("🗜️ Ukuran Hasil", f"{result_size} KB")

            st.success("✅ File berhasil diproses.")
            st.download_button("⬇️ Unduh File", data=result_io, file_name=download_name)

        progress.progress(100)
